"""Deterministic subcontract document render (SC-S3b) — the filled body TEXT → an EDITABLE .docx, and
the Schedule of Values → an EDITABLE .xlsx. NO AI, NO send (Invariant 1 generation side; capability-
gated). The operator's decision (2026-07-11): the deliverables are editable Office files (python-docx /
openpyxl), NOT flat PDF, so specific clauses / line values can be adjusted by hand before signature.

Layered on the SC-S3a render core: `render_subcontract_docx` calls
`subcontract_generate.render_body_text`, which runs ALL the correctness gates (SOV-sums-to-price,
Layer-A legal-review, strict token fill) BEFORE any bytes are produced — a bad money figure, an
un-cleared terms version, or an unfilled contract blank RAISES and no document is written. The .xlsx
render independently re-runs the SOV guard.

DETERMINISTIC — byte-identical output for a fixed record, which is what §47 version-on-conflict Box
filing relies on (SC-S3c skips a redundant upload when the recompiled bytes are unchanged, mirroring
the PO PDF's reportlab `invariant=1` CreationDate pin). No clock reads. `_normalize_ooxml_clock` pins BOTH
wall-clock sources an OOXML package carries to the record's agreement date (`agreement_ymd`, never
`datetime.now()`):
  1. `docProps/core.xml`'s `<dcterms:created>`/`<dcterms:modified>` CONTENT (openpyxl overwrites
     `modified` with `now()` at save time, ignoring the `wb.properties` override); and
  2. every ZIP member's local-header `date_time` (openpyxl stamps these from wall-clock too).
Both formats are normalized so the byte-determinism guarantee is uniform (python-docx alone is already
stable, but is normalized identically).
"""
from __future__ import annotations

import io
import re
import zipfile
from datetime import date, datetime
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font
from openpyxl.utils import get_column_letter

from subcontracts import exhibit, governing_law, money, subcontract_generate, terms

# A top-level article heading is `<n>.<TAB>TITLE:` — a SINGLE number, a dot, then a tab (the corpus
# body's structure). Sub-clauses are `<n>.<m>` (a second number right after the dot), so this pattern
# cleanly discriminates the 27 article headers from their sub-clauses without a hand-kept list.
_ARTICLE_HEADING_RE = re.compile(r"^\d+\.\t")
# Exhibit A uses a DIFFERENT heading convention than the subcontract body: `Article <ROMAN>:` /
# `ARTICLE <ROMAN>:` lines (e.g. "Article I: General", "ARTICLE V: PAYMENT REQUIREMENTS"). The Exhibit
# title ("Exhibit A: Scope of Work") is handled separately as the centered document title, so it never
# reaches this pattern. Trade Article II bodies use clause ids (e.g. "C2.0", "1.") — deliberately NOT
# matched, they render as justified paragraphs.
_EXHIBIT_HEADING_RE = re.compile(r"^Article\s+[IVXLCDM]+\b", re.IGNORECASE)
# The document title line (verbatim first content line of the body).
_TITLE_TEXT = "SUBCONTRACT AGREEMENT"


class SubcontractDocxError(Exception):
    """The subcontract package can't be rendered (bad record / SOV mismatch / gate failure). The daemon
    fences this to the Review Queue and NEVER files a package whose numbers/clauses don't re-derive."""


def _agreement_datetime(subcontract: dict[str, Any]) -> datetime:
    """The record's agreement date as a naive datetime (midnight) — the deterministic core-property
    timestamp. Constructed from `agreement_ymd`, never a clock read; raises on a malformed value."""
    ymd = subcontract.get("agreement_ymd")
    if not (isinstance(ymd, (list, tuple)) and len(ymd) == 3):
        raise SubcontractDocxError("subcontract record missing agreement_ymd (year, month, day)")
    try:
        return datetime(int(ymd[0]), int(ymd[1]), int(ymd[2]))
    except (ValueError, TypeError) as exc:
        raise SubcontractDocxError(f"invalid agreement_ymd {ymd!r}: {exc}") from exc


_CORE_DT_RE = re.compile(rb"(<dcterms:(created|modified)[^>]*>)[^<]+(</dcterms:\2>)")


def _normalize_ooxml_clock(data: bytes, stamp: datetime) -> bytes:
    """Rebuild an OOXML (.docx/.xlsx) ZIP with BOTH wall-clock sources pinned to `stamp` (the agreement
    date), so a re-render of the same record is byte-identical — what §47 idempotent Box filing needs:

      1. every member's ZIP local-header `date_time` (openpyxl's `wb.save()` stamps these from
         wall-clock, independent of any property override); and
      2. `docProps/core.xml`'s `<dcterms:created>` / `<dcterms:modified>` CONTENT — openpyxl overwrites
         `modified` with `now()` at save time regardless of `wb.properties.modified`, so the property
         pin alone is insufficient; we rewrite the element text to the stamp.

    Member order + all other content are preserved. Deterministic: same input → same output (zlib
    deflate is deterministic for a fixed level)."""
    # ZIP dates are bounded to [1980, 2107]; clamp defensively (agreement dates are contemporary).
    dt = (min(2107, max(1980, stamp.year)), stamp.month, stamp.day, 0, 0, 0)
    iso = f"{stamp.year:04d}-{stamp.month:02d}-{stamp.day:02d}T00:00:00Z".encode()
    out = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(data), "r") as src, \
            zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as dst:
        for item in src.infolist():
            content = src.read(item.filename)
            if item.filename == "docProps/core.xml":
                content = _CORE_DT_RE.sub(rb"\g<1>" + iso + rb"\g<3>", content)
            zi = zipfile.ZipInfo(item.filename, date_time=dt)
            zi.compress_type = item.compress_type
            zi.external_attr = item.external_attr
            zi.internal_attr = item.internal_attr
            zi.create_system = item.create_system
            dst.writestr(zi, content)
    return out.getvalue()


def render_subcontract_docx(
    subcontract: dict[str, Any],
    sov_lines: list[dict[str, Any]],
    *,
    terms_profile_id: str = "standard_subcontract",
    terms_version: str | None = None,
) -> bytes:
    """The filled 27-article subcontract body as an EDITABLE .docx (bytes). Runs the full SC-S3a gate
    chain (SOV guard → Layer-A legal gate → strict token fill) via `render_body_text` FIRST; only on a
    clean render are bytes produced. Article headers are bolded; every other clause is a justified
    paragraph carrying its source numbering verbatim (the operator edits from here)."""
    # Every gate the render core enforces (shape/SOV → SubcontractGenerateError, price-words →
    # MoneyError, Layer-A legal / sha / attach-kind → TermsError, unknown state → GoverningLawError)
    # is surfaced as this module's single fence type, so the daemon catches one exception and NEVER
    # files a wrong contract. A non-gate bug (AttributeError etc.) deliberately propagates uncaught.
    try:
        body = subcontract_generate.render_body_text(
            subcontract, sov_lines, terms_profile_id=terms_profile_id, terms_version=terms_version
        )
    except (
        subcontract_generate.SubcontractGenerateError,
        money.MoneyError,
        terms.TermsError,
        governing_law.GoverningLawError,
    ) as exc:
        raise SubcontractDocxError(f"body render gate failed: {type(exc).__name__}: {exc}") from exc

    doc = Document()
    # Base style: a serif legal body at 11pt, tight spacing.
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)

    lines = body.split("\n")
    first_seen = False
    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            continue
        # The title line → centered heading, once.
        if not first_seen and line.strip() == _TITLE_TEXT:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(_TITLE_TEXT)
            run.bold = True
            run.font.size = Pt(14)
            first_seen = True
            continue
        first_seen = True
        if _ARTICLE_HEADING_RE.match(line):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
        else:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Pin core-property timestamps to the agreement date (deterministic; no clock read).
    stamp = _agreement_datetime(subcontract)
    doc.core_properties.created = stamp
    doc.core_properties.modified = stamp
    doc.core_properties.author = str(subcontract.get("contractor_entity") or "Evergreen Renewables LLC")
    doc.core_properties.title = f"Subcontract — {subcontract.get('project_name', '')}".strip(" —")

    buf = io.BytesIO()
    doc.save(buf)
    return _normalize_ooxml_clock(buf.getvalue(), stamp)


def render_exhibit_a_docx(subcontract: dict[str, Any], contractor: dict[str, Any]) -> bytes:
    """Exhibit A (Scope of Work) as an EDITABLE .docx (bytes). Loads the FIXED, sha256-pinned skeleton
    (``exhibit.load_skeleton`` — Article I General + a ``{{article_ii}}`` marker + Articles III/IV/V/VI)
    and fills its eight tokens from the record. ``{{article_ii}}`` is the OPERATOR-authored
    ``exhibit_a_work_text`` when non-blank, else the trade's standard Article II "The Work" body
    (``exhibit.load_trade_art2`` — the trade-template fallback).

    Token fill is STRICT (``exhibit.substitute_tokens`` raises on a blank/missing contract token — a
    subcontract must never render with an unfilled blank). Any ``exhibit.ExhibitError`` (unknown trade,
    sha drift, unfilled token) is fenced as ``SubcontractDocxError`` so the daemon catches one type and
    never files a wrong Exhibit. Deterministic like the body render: the OOXML clock is pinned to the
    record's agreement date (``_normalize_ooxml_clock``), so a re-render of the same record is
    byte-identical — the §47 idempotent-Box-filing guarantee. The Exhibit title is a centered heading;
    every ``Article <ROMAN>:`` line is bolded; every other line is a justified paragraph the operator
    edits from."""
    try:
        skeleton = exhibit.load_skeleton()
        work_text = subcontract.get("exhibit_a_work_text")
        if isinstance(work_text, str) and work_text.strip():
            article_ii = work_text
        else:
            article_ii = exhibit.load_trade_art2(str(subcontract["trade"]))
        values = {
            "subcontractor_entity": str(subcontract["subcontractor_entity"]),
            "contractor_entity": str(contractor["entity"]),
            "trade": str(subcontract["trade"]),
            "project_name": str(subcontract["project_name"]),
            # site_address / completion_date are OPTIONAL D1 columns (0050 DEFAULT ''). A blank one
            # gets an editable bracketed placeholder rather than fencing the WHOLE package (the
            # operator fills it in the editable .docx) — only the required PARTY/trade tokens are
            # strict. owner_entity is required by the body's _REQUIRED_FIELDS, so it is always present.
            "site_address": str(subcontract.get("site_address") or "").strip() or "[site address]",
            "owner_entity": str(subcontract["owner_entity"]),
            "completion_date": str(subcontract.get("completion_date") or "").strip()
            or "[completion date to be confirmed]",
            "article_ii": article_ii,
        }
        text = exhibit.substitute_tokens(skeleton, values)
    except exhibit.ExhibitError as exc:
        raise SubcontractDocxError(
            f"exhibit render gate failed: {type(exc).__name__}: {exc}"
        ) from exc

    doc = Document()
    # Base style: a serif legal body at 11pt (mirrors the subcontract body render).
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)

    first_line = True
    for raw in text.split("\n"):
        line = raw.rstrip()
        if not line.strip():
            continue
        # The very first content line is the Exhibit title → centered heading, once.
        if first_line:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(14)
            first_line = False
            continue
        if _EXHIBIT_HEADING_RE.match(line):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
        else:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Pin core-property timestamps to the agreement date (deterministic; no clock read).
    stamp = _agreement_datetime(subcontract)
    doc.core_properties.created = stamp
    doc.core_properties.modified = stamp
    doc.core_properties.author = str(
        contractor.get("entity") or subcontract.get("contractor_entity") or "Evergreen Renewables LLC"
    )
    doc.core_properties.title = f"Exhibit A — {subcontract.get('project_name', '')}".strip(" —")

    buf = io.BytesIO()
    doc.save(buf)
    return _normalize_ooxml_clock(buf.getvalue(), stamp)


def _dollars(cents: int) -> float:
    """Integer cents → a dollars float for a SPREADSHEET CELL only (presentation, not the money path —
    the SOV guard reconciles in integer cents). Kept isolated so no float touches the correctness gate."""
    return cents / 100.0


def render_sov_xlsx(subcontract: dict[str, Any], sov_lines: list[dict[str, Any]]) -> bytes:
    """The Annex C Schedule of Values as an EDITABLE .xlsx (bytes). Independently re-runs the
    SOV-sums-to-price guard (a mismatch RAISES — never renders a workbook whose lines don't reconcile
    to §2.1). One row per SOV line (Scope | Quantity | Unit Price | Extended Value) + a TOTAL row that
    equals the Contract Price. Currency cells are numbers with a $ format so the operator can adjust."""
    price_cents = subcontract.get("contract_price_cents")
    if not isinstance(price_cents, int) or isinstance(price_cents, bool) or price_cents < 0:
        raise SubcontractDocxError(f"contract_price_cents must be a non-negative integer (got {price_cents!r})")
    problems = money.sov_mismatches(price_cents, sov_lines)
    if problems:
        raise SubcontractDocxError("SOV does not reconcile to the Contract Price: " + "; ".join(problems))

    wb = Workbook()
    ws = wb.active
    ws.title = "Schedule of Values"
    money_fmt = '"$"#,##0.00'
    bold = Font(bold=True)

    # Header block.
    ws["A1"] = "SCHEDULE OF VALUES (Annex C)"
    ws["A1"].font = Font(bold=True, size=14)
    ws["A2"] = "Project:"
    ws["B2"] = str(subcontract.get("project_name") or "")
    ws["A3"] = "Subcontractor:"
    ws["B3"] = str(subcontract.get("subcontractor_entity") or "")
    ws["A4"] = "Contract Price:"
    ws["B4"] = _dollars(price_cents)
    ws["B4"].number_format = money_fmt
    for cell in ("A2", "A3", "A4"):
        ws[cell].font = bold

    # Table header.
    header_row = 6
    headers = ["#", "Scope of Work", "Quantity", "Unit Price", "Extended Value"]
    for col, text in enumerate(headers, start=1):
        c = ws.cell(row=header_row, column=col, value=text)
        c.font = bold
        c.alignment = Alignment(horizontal="center")

    total = 0
    row = header_row + 1
    for i, line in enumerate(sov_lines, start=1):
        scope = str(line.get("description") or line.get("scope") or "Scope of Work")
        unit = line.get("unit_price_cents")
        stated = line.get("extended_cents")
        if unit is None:
            ext_cents = int(stated) if isinstance(stated, int) and not isinstance(stated, bool) else 0
            qty_val: Any = ""
            unit_val: Any = ""
        else:
            qty = line.get("qty", 1)
            ext_cents = money.sov_extended_cents(float(qty), int(unit))
            qty_val = qty
            unit_val = _dollars(int(unit))
        total += ext_cents
        ws.cell(row=row, column=1, value=i)
        ws.cell(row=row, column=2, value=scope)
        ws.cell(row=row, column=3, value=qty_val)
        uc = ws.cell(row=row, column=4, value=unit_val)
        if unit_val != "":
            uc.number_format = money_fmt
        ec = ws.cell(row=row, column=5, value=_dollars(ext_cents))
        ec.number_format = money_fmt
        row += 1

    # TOTAL row (equals the Contract Price — the guard above proved total == price).
    tlabel = ws.cell(row=row, column=4, value="TOTAL")
    tlabel.font = bold
    tcell = ws.cell(row=row, column=5, value=_dollars(total))
    tcell.font = bold
    tcell.number_format = money_fmt

    # Column widths for legibility.
    widths = {1: 6, 2: 48, 3: 12, 4: 16, 5: 18}
    for col, w in widths.items():
        ws.column_dimensions[get_column_letter(col)].width = w

    # Deterministic core-property timestamps (openpyxl defaults these to datetime.now()).
    stamp = _agreement_datetime(subcontract)
    wb.properties.created = stamp
    wb.properties.modified = stamp
    wb.properties.creator = str(subcontract.get("contractor_entity") or "Evergreen Renewables LLC")
    wb.properties.title = f"Schedule of Values — {subcontract.get('project_name', '')}".strip(" —")

    buf = io.BytesIO()
    wb.save(buf)
    return _normalize_ooxml_clock(buf.getvalue(), stamp)


def render_package(
    subcontract: dict[str, Any],
    sov_lines: list[dict[str, Any]],
    *,
    terms_profile_id: str = "standard_subcontract",
    terms_version: str | None = None,
) -> dict[str, bytes]:
    """The editable subcontract package (SC-S3b core): the Subcontract body .docx + the Exhibit A
    Scope-of-Work .docx + the Annex C Schedule-of-Values .xlsx, keyed by filename. Every gate runs; any
    failure raises SubcontractDocxError (the daemon fences, never files a partial/wrong package). The
    Exhibit A render loads the sha-pinned trade-template config and fills the operator's Article II (or
    the trade fallback)."""
    docx_bytes = render_subcontract_docx(
        subcontract, sov_lines, terms_profile_id=terms_profile_id, terms_version=terms_version
    )
    exhibit_bytes = render_exhibit_a_docx(subcontract, terms.load_contractor_config())
    xlsx_bytes = render_sov_xlsx(subcontract, sov_lines)
    return {
        "Subcontract.docx": docx_bytes,
        "Exhibit A.docx": exhibit_bytes,
        "Annex C - Schedule of Values.xlsx": xlsx_bytes,
    }


def zip_package(package: dict[str, bytes], stamp: datetime | date) -> bytes:
    """Combine a rendered ``package`` (the keyed files from ``render_package``) into ONE
    deterministic ZIP — the SC-S4 SEND artifact (2026-07-15 operator decision: the
    subcontractor receives the whole signable package as a single ``Subcontract Package.zip``,
    so the shared single-attachment send engine sends it unchanged).

    DETERMINISTIC — byte-identical for a fixed record, so §47 version-on-conflict Box filing
    still skips a redundant re-upload (the same guarantee ``render_package``'s members carry).
    Three sources of ZIP non-determinism are pinned: (1) every member's local-header
    ``date_time`` is set to the record's agreement date (never ``datetime.now()``), mirroring
    ``_normalize_ooxml_clock``; (2) members are written in sorted-name order; (3) a fixed
    ``ZIP_DEFLATED`` level (zlib deflate is deterministic for a fixed level). The inner
    ``.docx``/``.xlsx`` bytes are already deterministic (``render_package`` pinned their OOXML
    clocks), so the resulting ZIP is stable input-for-input."""
    dt = (min(2107, max(1980, stamp.year)), stamp.month, stamp.day, 0, 0, 0)
    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zf:
        for name in sorted(package):
            zi = zipfile.ZipInfo(name, date_time=dt)
            zi.compress_type = zipfile.ZIP_DEFLATED
            zi.external_attr = 0o600 << 16  # -rw------- (deterministic member perms)
            zf.writestr(zi, package[name])
    return out.getvalue()
