"""Tests for subcontracts/subcontract_docx.py (SC-S3b) — the editable .docx / .xlsx render layered on
the S3a text core. Renders are parsed BACK (python-docx / openpyxl) to assert content, mirroring the
render-parity discipline of tests/test_form_pdf.py. Uses a tmp CLEARED terms fixture (never the live
seed's pending version — HOUSE_REFLEXES §5)."""
from __future__ import annotations

import hashlib
import io
import json

import pytest
from docx import Document
from openpyxl import load_workbook

from subcontracts import exhibit, terms
from subcontracts import subcontract_docx as sd
from subcontracts.subcontract_docx import SubcontractDocxError

_CONTRACTOR = {
    "entity": "Evergreen Renewables LLC", "signature_entity": "Evergreen Renewables LLC",
    "prime_contractor_default": "Evergreen Renewables of Virginia LLC",
    "address_lines": ["100 Spectrum"], "phone": "1", "config_version": 1,
}


def _record(**over):
    base = {
        "subcontractor_entity": "D.E.L. Electric OR, Inc.", "project_name": "Kendall Solar",
        "owner_entity": "Kendall Solar, LLC", "governing_law_state": "OR",
        "contract_price_cents": 27401850, "price_basis": "fixed", "agreement_ymd": [2026, 7, 11],
        # Exhibit A fields (D1 migration-0050 columns) — needed so render_package's Exhibit A render
        # (strict token fill) has non-blank values. Extra keys are harmless to the body/SOV renders.
        "trade": "Civil", "site_address": "123 Field Rd, County, OR",
        "completion_date": "December 31, 2026", "exhibit_a_work_text": "",
    }
    base.update(over)
    return base


def _seed_terms(tmp_path, legal_review="cleared"):
    # A representative body: title + preamble + an article heading (`N.<TAB>TITLE:`) + clauses + the
    # price/law tokens — exercises BOTH the bold-heading branch and the justified-paragraph branch.
    body = (
        b"SUBCONTRACT AGREEMENT\n\n"
        b"THIS AGREEMENT by and between {{contractor_entity}} and {{subcontractor_entity}} "
        b"for {{project_name}}.\n"
        b"1.\tSCOPE OF WORK:\n"
        b"1.1\tThe Subcontractor shall perform the Work described in Exhibit A.\n"
        b"2.\tCONTRACT PRICE:\n"
        b"2.1\tThe Contract Price is {{contract_price_clause}}, governed by the laws of "
        b"{{governing_law_state_name}}.\n"
    )
    tdir = tmp_path / "terms"
    tdir.mkdir()
    (tdir / "b_v1.md").write_bytes(body)
    (tdir / "manifest.json").write_text(json.dumps({
        "manifest_version": 1, "profiles": {"standard_subcontract": {
            "kind": "library", "current_version": "v1",
            "versions": {"v1": {"file": "b_v1.md", "sha256": hashlib.sha256(body).hexdigest(),
                                "tokens": ["contractor_entity", "subcontractor_entity", "project_name",
                                           "contract_price_clause", "governing_law_state_name"],
                                "legal_review": legal_review}}}}}), encoding="utf-8")
    cdir = tmp_path / "config"
    cdir.mkdir()
    (cdir / "contractor.json").write_text(json.dumps({**_CONTRACTOR}), encoding="utf-8")
    return tdir, cdir


def _cleared(tmp_path, monkeypatch, legal_review="cleared"):
    tdir, cdir = _seed_terms(tmp_path, legal_review=legal_review)
    monkeypatch.setattr(terms, "TERMS_DIR", tdir)
    monkeypatch.setattr(terms, "CONFIG_DIR", cdir)


_SOV = [{"description": "Solar electrical scope", "extended_cents": 27401850}]


# ── .docx render ─────────────────────────────────────────────────────────────


def test_render_subcontract_docx_fills_and_parses(tmp_path, monkeypatch):
    _cleared(tmp_path, monkeypatch)
    data = sd.render_subcontract_docx(_record(), _SOV)
    assert isinstance(data, bytes) and data[:2] == b"PK"  # a real zip/.docx
    doc = Document(io.BytesIO(data))
    texts = [p.text for p in doc.paragraphs]
    joined = "\n".join(texts)
    assert texts[0] == "SUBCONTRACT AGREEMENT"
    assert "Evergreen Renewables LLC and D.E.L. Electric OR, Inc." in joined
    assert "Two hundred seventy-four thousand eighteen dollars and fifty cents ($274,018.50)" in joined
    assert "the laws of the State of Oregon" in joined
    assert "{{" not in joined  # no unfilled tokens survived to the document


def test_render_subcontract_docx_bolds_article_headings(tmp_path, monkeypatch):
    _cleared(tmp_path, monkeypatch)
    doc = Document(io.BytesIO(sd.render_subcontract_docx(_record(), _SOV)))
    # The `1.<TAB>SCOPE OF WORK:` and `2.<TAB>CONTRACT PRICE:` lines are bold; a normal clause is not.
    bold_texts = [p.text for p in doc.paragraphs if p.runs and p.runs[0].bold]
    assert any(t.startswith("1.") and "SCOPE OF WORK" in t for t in bold_texts)
    assert any(t.startswith("2.") and "CONTRACT PRICE" in t for t in bold_texts)
    # A sub-clause (1.1 …) is NOT a bold heading.
    assert not any(t.startswith("1.1") for t in bold_texts)


def test_render_subcontract_docx_fences_pending_body(tmp_path, monkeypatch):
    _cleared(tmp_path, monkeypatch, legal_review="pending")
    with pytest.raises(SubcontractDocxError, match="gate failed"):
        sd.render_subcontract_docx(_record(), _SOV)


def test_render_subcontract_docx_fences_unknown_state(tmp_path, monkeypatch):
    _cleared(tmp_path, monkeypatch)
    with pytest.raises(SubcontractDocxError, match="gate failed"):
        sd.render_subcontract_docx(_record(governing_law_state="ZZ"), _SOV)


def test_render_subcontract_docx_fences_sov_mismatch(tmp_path, monkeypatch):
    _cleared(tmp_path, monkeypatch)
    with pytest.raises(SubcontractDocxError, match="gate failed"):
        sd.render_subcontract_docx(_record(), [{"extended_cents": 9999}])  # != price


# ── Exhibit A .docx render ───────────────────────────────────────────────────

# Exhibit A takes the contractor as a param (not the terms config dir) — no terms monkeypatch needed;
# it loads the LIVE sha-pinned exhibit config. Corpus markers are DERIVED from exhibit.load_trade_art2
# (never pinned as literals), mirroring test_subcontract_exhibit.py's HOUSE_REFLEXES §5 discipline.


def _exhibit_record(**over):
    base = {
        "subcontractor_entity": "D.E.L. Electric OR, Inc.", "contractor_entity": "Evergreen Renewables LLC",
        "trade": "Civil", "project_name": "Kendall Solar", "owner_entity": "Kendall Solar, LLC",
        "site_address": "123 Field Rd, County, OR", "completion_date": "December 31, 2026",
        "agreement_ymd": [2026, 7, 11],
    }
    base.update(over)
    return base


def _exhibit_text(data: bytes) -> str:
    return "\n".join(p.text for p in Document(io.BytesIO(data)).paragraphs)


def _civil_marker() -> str:
    # A distinctive clause line from the Civil trade template, derived from source (not a literal pin).
    return exhibit.load_trade_art2("Civil").strip().splitlines()[1][:30]


def test_render_exhibit_a_docx_fills_tokens_and_trade_fallback():
    data = sd.render_exhibit_a_docx(_exhibit_record(), _CONTRACTOR)
    assert isinstance(data, bytes) and data[:2] == b"PK"
    joined = _exhibit_text(data)
    # Skeleton tokens all filled.
    assert "D.E.L. Electric OR, Inc." in joined
    assert "Evergreen Renewables LLC" in joined
    assert "Kendall Solar" in joined
    assert "Kendall Solar, LLC" in joined
    assert "123 Field Rd, County, OR" in joined
    assert "December 31, 2026" in joined
    assert "{{" not in joined  # no unfilled token survived to the document
    # exhibit_a_work_text blank → the trade's standard Article II body was substituted.
    assert _civil_marker() in joined


def test_render_exhibit_a_docx_uses_operator_work_text_when_provided():
    sentinel = "Bespoke operator-authored scope clause ZZZ-42."
    data = sd.render_exhibit_a_docx(_exhibit_record(exhibit_a_work_text=sentinel), _CONTRACTOR)
    joined = _exhibit_text(data)
    assert sentinel in joined                 # the operator's Article II text is used
    assert _civil_marker() not in joined      # the trade-template fallback was NOT used


def test_render_exhibit_a_docx_bolds_article_headings():
    doc = Document(io.BytesIO(sd.render_exhibit_a_docx(_exhibit_record(), _CONTRACTOR)))
    bold_texts = [p.text for p in doc.paragraphs if p.runs and p.runs[0].bold]
    # At least one `Article <ROMAN>:` / `ARTICLE <ROMAN>:` skeleton heading is bold.
    assert any(t.lower().startswith("article ") for t in bold_texts)


def test_render_exhibit_a_docx_fences_unknown_trade():
    with pytest.raises(SubcontractDocxError, match="exhibit render gate failed"):
        sd.render_exhibit_a_docx(_exhibit_record(trade="Underwater Basket Weaving"), _CONTRACTOR)


def test_render_exhibit_a_docx_fences_blank_required_token():
    # A blank REQUIRED party token (project_name) is a strict-fill failure, fenced as
    # SubcontractDocxError. (site_address / completion_date are optional-tolerant — see below.)
    with pytest.raises(SubcontractDocxError, match="exhibit render gate failed"):
        sd.render_exhibit_a_docx(_exhibit_record(project_name=""), _CONTRACTOR)


def test_render_exhibit_a_docx_tolerates_blank_optional_fields():
    # A blank optional site_address / completion_date renders with an editable bracketed placeholder
    # (never fences the whole package — the operator fills it in the .docx).
    data = sd.render_exhibit_a_docx(_exhibit_record(site_address="", completion_date=""), _CONTRACTOR)
    text = "\n".join(p.text for p in Document(io.BytesIO(data)).paragraphs)
    assert "[site address]" in text and "[completion date to be confirmed]" in text


def test_render_exhibit_a_docx_byte_deterministic():
    d1 = sd.render_exhibit_a_docx(_exhibit_record(), _CONTRACTOR)
    d2 = sd.render_exhibit_a_docx(_exhibit_record(), _CONTRACTOR)  # same record, re-render
    assert d1 == d2, "Exhibit A .docx render is not byte-deterministic"
    # And the OOXML clock is pinned to the agreement date, not a wall-clock read.
    import zipfile
    core = zipfile.ZipFile(io.BytesIO(d1)).read("docProps/core.xml").decode()
    assert "2026-07-11T00:00:00Z</dcterms:modified>" in core, core


# ── .xlsx SOV render ─────────────────────────────────────────────────────────


def test_render_sov_xlsx_table_and_total(tmp_path, monkeypatch):
    data = sd.render_sov_xlsx(_record(), _SOV)
    assert isinstance(data, bytes) and data[:2] == b"PK"
    ws = load_workbook(io.BytesIO(data)).active
    assert ws["A1"].value == "SCHEDULE OF VALUES (Annex C)"
    assert ws["B4"].value == pytest.approx(274018.50)  # contract price dollars
    # header row at 6, one data row at 7, TOTAL row at 8
    assert [ws.cell(row=6, column=c).value for c in range(1, 6)] == \
        ["#", "Scope of Work", "Quantity", "Unit Price", "Extended Value"]
    assert ws.cell(row=7, column=2).value == "Solar electrical scope"
    assert ws.cell(row=7, column=5).value == pytest.approx(274018.50)
    assert ws.cell(row=8, column=4).value == "TOTAL"
    assert ws.cell(row=8, column=5).value == pytest.approx(274018.50)


def test_render_sov_xlsx_multiline_unit_priced(tmp_path, monkeypatch):
    sov = [
        {"description": "Panels", "qty": 2, "unit_price_cents": 5000000, "extended_cents": 10000000},
        {"description": "Labor", "qty": 1, "unit_price_cents": 10000000, "extended_cents": 10000000},
    ]
    ws = load_workbook(io.BytesIO(sd.render_sov_xlsx(_record(contract_price_cents=20000000), sov))).active
    assert ws.cell(row=7, column=3).value == 2            # qty
    assert ws.cell(row=7, column=4).value == pytest.approx(50000.00)  # unit price
    assert ws.cell(row=8, column=5).value == pytest.approx(100000.00)  # 2nd line extended
    # TOTAL row is now row 9 (2 data rows)
    assert ws.cell(row=9, column=5).value == pytest.approx(200000.00)


def test_render_sov_xlsx_fences_on_mismatch(tmp_path, monkeypatch):
    with pytest.raises(SubcontractDocxError, match="reconcile"):
        sd.render_sov_xlsx(_record(), [{"extended_cents": 9999}])  # != price
    with pytest.raises(SubcontractDocxError, match="non-negative integer"):
        sd.render_sov_xlsx(_record(contract_price_cents=-1), _SOV)


# ── package + determinism ────────────────────────────────────────────────────


def test_render_package_returns_three_artifacts(tmp_path, monkeypatch):
    _cleared(tmp_path, monkeypatch)
    pkg = sd.render_package(_record(), _SOV)
    assert set(pkg) == {"Subcontract.docx", "Exhibit A.docx", "Annex C - Schedule of Values.xlsx"}
    assert all(isinstance(v, bytes) and v[:2] == b"PK" and len(v) > 500 for v in pkg.values())


def test_deterministic_core_property_timestamp(tmp_path, monkeypatch):
    _cleared(tmp_path, monkeypatch)
    d1 = Document(io.BytesIO(sd.render_subcontract_docx(_record(), _SOV)))
    d2 = Document(io.BytesIO(sd.render_subcontract_docx(_record(), _SOV)))  # same fixture, re-render
    # Pinned to the agreement date, not a clock read → identical across renders.
    assert d1.core_properties.created == d2.core_properties.created
    assert str(d1.core_properties.created).startswith("2026-07-11")
    assert [p.text for p in d1.paragraphs] == [p.text for p in d2.paragraphs]


def test_byte_identical_across_renders_both_formats(tmp_path, monkeypatch):
    """§47-readiness: a re-render of the SAME record must be BYTE-identical for BOTH formats — this is
    what lets SC-S3c skip a redundant Box upload on an unchanged recompile. The .xlsx is the one that
    regressed (openpyxl stamps zip members from wall-clock); _normalize_ooxml_clock fixes it. Guards
    against silent regression of that guarantee."""
    _cleared(tmp_path, monkeypatch)
    docx1 = sd.render_subcontract_docx(_record(), _SOV)
    docx2 = sd.render_subcontract_docx(_record(), _SOV)
    xlsx1 = sd.render_sov_xlsx(_record(), _SOV)
    xlsx2 = sd.render_sov_xlsx(_record(), _SOV)
    assert docx1 == docx2, "subcontract .docx render is not byte-deterministic"
    assert xlsx1 == xlsx2, "SOV .xlsx render is not byte-deterministic"
    # And both still open as valid OOXML after the clock normalization.
    assert Document(io.BytesIO(docx1)).paragraphs[0].text == "SUBCONTRACT AGREEMENT"
    assert load_workbook(io.BytesIO(xlsx1)).active["A1"].value == "SCHEDULE OF VALUES (Annex C)"
    # Directly assert NO wall-clock leak (equality alone can false-pass if two renders hit the same
    # second): docProps/core.xml's dcterms:modified must be the PINNED agreement date, not a clock time.
    import zipfile
    for data in (docx1, xlsx1):
        core = zipfile.ZipFile(io.BytesIO(data)).read("docProps/core.xml").decode()
        assert "2026-07-11T00:00:00Z</dcterms:modified>" in core, core
